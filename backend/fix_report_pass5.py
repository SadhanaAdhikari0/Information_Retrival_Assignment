"""
fix_report_pass5.py — Corrects an off-by-one bug from fix_report.py's table
update code: every table in this report has an extra blank leading row
(row 0) before the real header (row 1) — an artifact of the original
generator (doc.add_table(rows=1, ...) followed by add_table_row() for the
header, never removing the initial blank row). fix_report.py's positional
`t.rows[1:]` zips therefore overwrote the HEADER row instead of the first
data row, shifting every value up by one and leaving one stale row
dangling at the bottom. This rebuilds tables 2, 3, 5, 6 correctly:
row 0 = blank (left as-is), row 1 = header (restored), rows 2.. = data.
"""
import os
import re
from docx import Document
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), ".env"))
from pymongo import MongoClient

db = MongoClient(os.environ["MONGODB_URI"])["vertical_search_engine"]
REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")

import sys
sys.path.insert(0, os.path.dirname(__file__))
from app import search_documents
from rss_collector import classify_text

model_doc = db.news_model_runs.find_one({}, sort=[("trained_at", -1)]) or {}
conf_matrix = model_doc.get("confusion_matrix")
conf_labels = model_doc.get("confusion_labels", ["Economics", "Entertainment", "Politics"])


def set_paragraph_text(paragraph, new_text, bold=None):
    if not paragraph.runs:
        run = paragraph.add_run(new_text)
    else:
        paragraph.runs[0].text = new_text
        run = paragraph.runs[0]
        for r in paragraph.runs[1:]:
            r.text = ""
    if bold is not None:
        run.bold = bold


def set_row(row, values, bold=False):
    for cell, val in zip(row.cells, values):
        set_paragraph_text(cell.paragraphs[0], str(val), bold=bold)


doc = Document(REPORT_PATH)
tables = doc.tables

# ---- Table[2]: live search evaluation (T1..T6) ----
t = tables[2]
set_row(t.rows[1], ["Test", "Query", "Results", "Top Result", "Top Score"], bold=True)
TEST_QUERIES = [
    ("T1", "mental health"),
    ("T2", "Deborah Lycett"),
    ("T3", "Celine Brookes-Smith"),
    ("T4", "nursing social care intervention"),
    ("T5", "healthcare community transformation"),
    ("T6", "machine learning quantum blockchain xyz"),
]
for row, (label, q) in zip(t.rows[2:], TEST_QUERIES):
    results, total = search_documents(q, page=1)
    top = results[0] if results else None
    set_row(row, [label, q, total,
                  (top["title"][:70] if top else "(no results)"),
                  (f"{top['score']:.4f}" if top else "N/A")])
print("Rebuilt Table[2].")

# ---- Table[3]: lexical-overlap precision proxy ----
t = tables[3]
set_row(t.rows[1], ["Query", "Relevant Retrieved", "Total Retrieved", "Precision@K"], bold=True)
proxy_queries = ["mental health", "Deborah Lycett", "nursing social care intervention"]
for row, q in zip(t.rows[2:], proxy_queries):
    results, total = search_documents(q, page=1)
    top10 = results[:10]
    q_words = {w.lower() for w in re.findall(r"[a-zA-Z]+", q) if len(w) > 2}
    relevant = sum(1 for r in top10
                    if q_words & {w.lower() for w in re.findall(r"[a-zA-Z]+", r["title"])})
    prec = relevant / len(top10) if top10 else 0.0
    set_row(row, [q, relevant, len(top10), f"{prec:.2f}"])
print("Rebuilt Table[3].")

# ---- Table[5]: confusion matrix ----
t = tables[5]
set_row(t.rows[1], ["True \\ Predicted"] + conf_labels, bold=True)
if conf_matrix:
    for row, true_label, counts_row in zip(t.rows[2:], conf_labels, conf_matrix):
        set_row(row, [true_label] + list(counts_row))
print("Rebuilt Table[5].")

# ---- Table[6]: live classification tests ----
t = tables[6]
set_row(t.rows[1], ["Expected Category", "Test Input (excerpt)", "Predicted", "Confidence"], bold=True)
CLASSIFY_TESTS = [
    ("Economics", "The Federal Reserve raised interest rates again as GDP growth "
                   "slowed and inflation remained above target."),
    ("Entertainment", "The Marvel blockbuster broke box office records this "
                       "weekend, grossing over 300 million dollars worldwide."),
    ("Politics", "Parliament voted to approve the new immigration bill after "
                 "the Prime Minister addressed the House of Commons."),
]
for row, (expected, text) in zip(t.rows[2:], CLASSIFY_TESTS):
    r = classify_text(text)
    set_row(row, [expected, text[:70] + "…", r["category"], f"{r['confidence']:.1%}"])
print("Rebuilt Table[6].")

doc.save(REPORT_PATH)
print("Saved pass 5.")
