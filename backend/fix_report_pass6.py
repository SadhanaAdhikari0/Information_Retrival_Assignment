"""fix_report_pass6.py — final table-label honesty fixes."""
import os
from docx import Document

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")


def set_paragraph_text(paragraph, new_text, bold=None):
    if not paragraph.runs:
        run = paragraph.add_run(new_text)
    else:
        paragraph.runs[0].text = new_text
        run = paragraph.runs[0]
        for r in paragraph.runs[1:]:
            r.text = ""
    if bold is not None:
        run.bold = bold


def set_row(row, values, bold=False):
    for cell, val in zip(row.cells, values):
        set_paragraph_text(cell.paragraphs[0], str(val), bold=bold)


doc = Document(REPORT_PATH)
tables = doc.tables

# Table[4]: dataset distribution — relabel around the REAL official
# requirement (>=100 documents total; the coursework brief does not require
# 150/category — that figure only appears in an unofficial aspirational spec).
t = tables[4]
set_row(t.rows[1], ["Category", "Documents Collected", "Percentage", "Source Mix"], bold=True)
row_sources = {
    "Economics": "BBC RSS + Wikipedia",
    "Entertainment": "BBC RSS + Wikipedia",
    "Politics": "BBC RSS + Wikipedia",
    "Total": "✓ Meets official ≥100-document minimum",
}
for row in t.rows[2:]:
    label = row.cells[0].text.strip()
    if label in row_sources:
        set_paragraph_text(row.cells[3].paragraphs[0], row_sources[label])

# Table[7]: requirements validation — fix the ≥450 requirement text to the
# real official minimum, and point at the real evidence section.
t = tables[7]
for row in t.rows[1:]:
    if row.cells[1].text.strip() == "News dataset (≥450, 3 categories)":
        set_row(row, ["7", "News dataset (≥100 total, 3 categories)",
                       "rss_collector.py: live BBC RSS (feedparser) + Wikipedia extracts",
                       "Section 3.4, Table 3"])

doc.save(REPORT_PATH)
print("Saved pass 6.")
