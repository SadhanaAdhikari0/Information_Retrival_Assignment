"""swap_ui_images.py — Replaces the stale UI screenshots with the fresh,
real ones just captured from the live app via Selenium."""
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")
SCREENSHOTS = os.path.join(os.path.dirname(__file__), "..", "screenshots")

# (exact caption paragraph text -- caption comes AFTER the image, filename)
REPLACEMENTS = [
    ("Figure 1: IR Research Platform — Vertical Search Engine Home Page",
     "ui_home_search_page.png"),
    ("Figure 3: Search Results for Query 'mental health' — Cosine Similarity Scores",
     "ui_search_mental_health.png"),
    ("Figure 4: News Document Clustering — Overview Dashboard (198 Documents)",
     "ui_news_overview.png"),
    ("Figure 5: K-Means Cluster Visualisation — PCA 2D Projection of 198 Real Documents",
     "ui_news_clusters.png"),
    ("Figure 6: Document Classification Panel — User Interface",
     "ui_classify_panel.png"),
    ("Figure 7: Classification Result — Live Prediction from the Trained K-Means Model",
     "ui_classify_result.png"),
]


def paragraph_has_image(p):
    return bool(p._element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"))


def clear_paragraph(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


doc = Document(REPORT_PATH)

for caption_text, filename in REPLACEMENTS:
    caption_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == caption_text:
            caption_idx = i
            break
    if caption_idx is None:
        print(f"  [WARN] caption not found: {caption_text}")
        continue

    img_para = None
    for j in range(caption_idx - 1, max(caption_idx - 4, -1), -1):
        if paragraph_has_image(doc.paragraphs[j]):
            img_para = doc.paragraphs[j]
            break

    img_path = os.path.join(SCREENSHOTS, filename)
    if img_para is None:
        print(f"  [WARN] no image paragraph found before: {caption_text}")
        continue
    if not os.path.exists(img_path):
        print(f"  [WARN] screenshot missing: {img_path}")
        continue

    clear_paragraph(img_para)
    run = img_para.add_run()
    run.add_picture(img_path, width=Inches(6.0))
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(f"  Replaced: {caption_text[:60]}")

doc.save(REPORT_PATH)
print("Saved.")
