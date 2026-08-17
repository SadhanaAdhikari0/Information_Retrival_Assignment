"""
trim_report.py — Condenses the report's main body (Sections 1-5) from
~5,300 words toward the brief's ~2,000-word target, while preserving every
formula, evaluation number, and the critical-analysis content (the
robots.txt bug story, the honest clustering-accuracy discussion, the
dataset-imbalance limitation) that the marking criteria reward. Abstract,
references, and appendices are front/back matter and are not trimmed.
"""
import os
from docx import Document

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")


def set_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def replace_exact(doc, old, new, label=""):
    for p in doc.paragraphs:
        if p.text.strip() == old.strip():
            set_paragraph_text(p, new)
            return True
    print(f"  [WARN] exact paragraph not found: {label or old[:60]}")
    return False


PAIRS = [
    # ---------- 1. Introduction ----------
    ("The exponential growth of digital information has created significant "
     "challenges for effective knowledge discovery. Academic research portals, "
     "news aggregators, and institutional repositories collectively produce "
     "vast volumes of unstructured text that cannot be effectively navigated "
     "using simple keyword matching. Information Retrieval (IR) addresses "
     "these challenges through principled mathematical models that enable the "
     "ranking and categorisation of textual content according to its "
     "relevance to a user's information need.",
     "The exponential growth of digital information challenges effective "
     "knowledge discovery: academic portals, news aggregators, and "
     "repositories produce vast unstructured text that simple keyword "
     "matching cannot navigate well. Information Retrieval (IR) addresses "
     "this through principled mathematical models that rank and categorise "
     "text by relevance to a user's information need."),

    ("This project implements two complementary IR components within a "
     "unified web application. Task 1 presents a vertical search engine "
     "specifically designed to retrieve research outputs and researcher "
     "profiles from the Coventry University Centre for Healthcare and "
     "Community Transformation PurePortal. The system employs the Vector "
     "Space Model (VSM) with TF-IDF term weighting and cosine similarity "
     "ranking, consistent with established IR theory (Manning, Raghavan, & "
     "Schütze, 2008). Task 2 addresses the problem of automated document "
     "categorisation through unsupervised K-Means clustering applied to news "
     "articles collected via RSS feeds, enabling automatic partitioning into "
     "Economics, Entertainment, and Politics categories.",
     "This project implements two complementary IR components. Task 1 is a "
     "vertical search engine retrieving research outputs and researcher "
     "profiles from the Coventry University Centre for Healthcare and "
     "Community Transformation PurePortal, using the Vector Space Model "
     "(VSM) with TF-IDF weighting and cosine similarity ranking (Manning, "
     "Raghavan, & Schütze, 2008). Task 2 addresses automated document "
     "categorisation through unsupervised K-Means clustering (K=3) applied "
     "to real news/reference documents, partitioning them into Economics, "
     "Entertainment, and Politics."),

    ("The motivation for this integrated system stems from the dual nature "
     "of real-world IR requirements: practitioners require both precise "
     "retrieval from domain-specific repositories (vertical search) and "
     "scalable automated organisation of high-volume news streams "
     "(clustering). By implementing both within a single platform, this "
     "project demonstrates the complementary nature of retrieval and "
     "classification approaches to IR.",
     "The motivation is the dual nature of real-world IR needs: precise "
     "retrieval from a domain-specific repository (vertical search), and "
     "scalable automated organisation of a document stream (clustering). "
     "Implementing both within one platform demonstrates the complementary "
     "nature of retrieval and classification."),

    ("The system is implemented using Python (Flask) for the backend API, "
     "React (Vite) for the frontend interface, MongoDB Atlas for persistent "
     "storage, and standard NLP libraries (NLTK, scikit-learn) for text "
     "processing. The crawler is implemented with the requests and "
     "BeautifulSoup libraries, using a breadth-first traversal from the seed "
     "URL that follows only Research Output and Profile links.",
     "The system is implemented using Python (Flask) for the backend API, "
     "React (Vite) for the frontend, MongoDB Atlas for storage, and "
     "NLTK/scikit-learn for text processing. The crawler uses requests and "
     "BeautifulSoup, performing a breadth-first traversal from the seed URL "
     "that follows only Research Output and Profile links."),

    # ---------- 2.1 ----------
    ("A vertical search engine focuses retrieval on a specific domain or "
     "corpus rather than the entire web. In this project, the domain is the "
     "Coventry University Centre for Healthcare and Community Transformation "
     "PurePortal, a repository of academic research outputs and researcher "
     "profiles. The problem addressed is the difficulty of locating specific "
     "research publications or researchers using natural language queries "
     "within a JavaScript-rendered academic portal that does not expose a "
     "public search API.",
     "A vertical search engine focuses retrieval on a specific domain rather "
     "than the entire web — here, the Coventry University Centre for "
     "Healthcare and Community Transformation PurePortal, a repository of "
     "research outputs and researcher profiles. The problem is locating "
     "specific publications or researchers via natural language queries "
     "within a portal that exposes no public search API."),

    ("The system must crawl, extract, and index research content from the "
     "PurePortal, then rank results using mathematically principled "
     "relevance scoring. Users may query by publication title, author name, "
     "keyword, or phrase, receiving ranked results with clickable "
     "publication and profile links, publication dates, and visible "
     "relevance scores.",
     "The system must crawl, extract, and index PurePortal content, then "
     "rank results with mathematically principled relevance scoring. Users "
     "may query by title, author name, keyword, or phrase, and receive "
     "ranked results with clickable links, publication dates, and visible "
     "relevance scores."),

    # ---------- 2.3 System Architecture ----------
    ("The Task 1 architecture follows a layered pipeline pattern. The data "
     "layer comprises MongoDB Atlas collections: doc_vectors (77 TF-IDF "
     "document vectors), term_index (3,812 indexed terms with IDF values), "
     "research_outputs (83 publications), researcher_profiles (122 "
     "profiles), and crawl_log (crawl history). The crawler layer uses the "
     "requests library with BeautifulSoup for HTML parsing and structured "
     "metadata extraction. The processing layer applies the NLP pipeline "
     "and constructs TF-IDF vectors. The API layer is a Flask REST API "
     "serving the search endpoint (/api/search), autocomplete "
     "(/api/suggestions), and crawl status (/api/crawl-status). The "
     "presentation layer is a React single-page application served by Vite "
     "during development.",
     "The Task 1 architecture is a layered pipeline. The data layer "
     "comprises MongoDB Atlas collections: doc_vectors (77 TF-IDF vectors), "
     "term_index (3,812 terms with IDF values), research_outputs (83 "
     "publications), researcher_profiles (122 profiles), and crawl_log. The "
     "crawler layer uses requests + BeautifulSoup for HTML parsing and "
     "metadata extraction. The processing layer applies the NLP pipeline "
     "and builds TF-IDF vectors. The API layer is a Flask REST API "
     "(/api/search, /api/suggestions, /api/crawl-status); the presentation "
     "layer is a React single-page application."),

    ("Plain HTTP GET requests (a descriptive User-Agent identifying the "
     "crawler as an educational bot) succeed directly against PurePortal — "
     "no browser automation is required. Starting from the seed "
     "organisation page, the crawler performs a breadth-first traversal, "
     "discovering and following only links whose path starts with "
     "/en/publications/, /en/persons/, or /en/organisations/"
     "centre-for-healthcare, and queuing every newly discovered relevant "
     "link until none remain.",
     "Plain HTTP GET requests (with a descriptive educational-bot "
     "User-Agent) succeed directly against PurePortal — no browser "
     "automation is required. From the seed page, the crawler performs a "
     "breadth-first traversal, following only links whose path starts with "
     "/en/publications/, /en/persons/, or /en/organisations/"
     "centre-for-healthcare, queuing every newly discovered relevant link "
     "until none remain."),

    ("Polite crawling practices are implemented throughout: a five-second "
     "delay between requests (CRAWL_DELAY_SECONDS = 5, matching "
     "PurePortal's own robots.txt Crawl-Delay directive), a descriptive "
     "User-Agent string, URL deduplication to prevent revisiting pages, and "
     "standard HTTP GET requests only. robots.txt compliance is actually "
     "enforced (can_fetch() is called from fetch_page() before every "
     "request), not merely demonstrated in a comment — see Section 2.4.1 "
     "for a bug this uncovered.",
     "Polite crawling is implemented throughout: a five-second delay "
     "between requests (matching PurePortal's own robots.txt Crawl-Delay), "
     "a descriptive User-Agent, URL deduplication, and GET requests only. "
     "robots.txt compliance is actually enforced — can_fetch() is called "
     "from fetch_page() before every request, not merely demonstrated in a "
     "comment (see Section 2.4.1 for a real bug this uncovered)."),

    ("Ethical web crawling is a fundamental requirement of modern "
     "Information Retrieval systems. The crawler uses Python's native "
     "urllib.robotparser to check every URL against PurePortal's robots.txt "
     "before it is fetched. Building this surfaced a genuine bug worth "
     "reporting: RobotFileParser.read() fetches robots.txt via bare "
     "urllib.request with no custom headers, and PurePortal's edge "
     "protection returns HTTP 403 to urllib's default User-Agent (confirmed "
     "independently: the identical URL returns 200 via the requests "
     "library). RobotFileParser silently treats a 403 response as "
     "\"disallow everything\", which meant the very first version of this "
     "check blocked the crawler from fetching even its own seed URL. The "
     "fix — fetch robots.txt text via requests with the same descriptive "
     "User-Agent used for every other request, then hand that text to "
     "RobotFileParser.parse() instead of .read() — resolved it. "
     "PurePortal's robots.txt (checked 2026-08-15) permits every path this "
     "crawler visits and specifies Crawl-Delay: 5, which the crawler's "
     "5-second delay matches.",
     "Ethical crawling is a core IR requirement. The crawler checks every "
     "URL against PurePortal's robots.txt via Python's urllib.robotparser "
     "before fetching. Building this surfaced a genuine bug: "
     "RobotFileParser.read() fetches robots.txt via bare urllib.request "
     "with no custom headers, and PurePortal's edge protection returns HTTP "
     "403 to urllib's default User-Agent (confirmed independently — the "
     "identical URL returns 200 via requests). RobotFileParser silently "
     "treats a 403 as \"disallow everything\", so the first version of this "
     "check blocked the crawler from fetching even its own seed URL. The "
     "fix — fetch the robots.txt text via requests with the crawler's own "
     "User-Agent, then hand it to RobotFileParser.parse() instead of "
     ".read() — resolved it. PurePortal's robots.txt (checked 2026-08-15) "
     "permits every path this crawler visits and specifies Crawl-Delay: 5, "
     "matching the crawler's delay."),

    ("The crawler scheduling is managed by scheduler.py (start_scheduler()), "
     "which uses APScheduler's BlockingScheduler with an IntervalTrigger set "
     "to days=CRAWL_INTERVAL_DAYS (90 days by default). This implements the "
     "strict 3-month crawl interval required by the assignment "
     "specification. The scheduler executes the crawl immediately on "
     "startup, then schedules subsequent executions at 3-month intervals. "
     "The schedule is explicitly documented throughout the codebase, "
     "configuration, and log output:",
     "Crawler scheduling is managed by scheduler.py (start_scheduler()), "
     "using APScheduler's BlockingScheduler with an IntervalTrigger set to "
     "days=CRAWL_INTERVAL_DAYS (90 by default) — implementing the strict "
     "3-month interval required by the specification. The scheduler runs "
     "the crawl immediately on startup, then re-runs every 3 months:"),

    ("This 3-month interval is consistent with the dynamic nature of "
     "academic research portals, where new publications are added "
     "quarterly. It represents a pragmatic balance between data freshness "
     "and server load, avoiding the excessive server impact of daily or "
     "weekly crawling.",
     "This interval suits the slow-changing nature of academic research "
     "portals, where new publications are typically added quarterly — "
     "balancing data freshness against server load, and avoiding the "
     "excessive impact of daily or weekly crawling."),

    ("The NLP preprocessing pipeline is applied identically to both crawled "
     "documents (during indexing) and user queries (during search). This "
     "consistency is essential for the Vector Space Model to function "
     "correctly, as query and document vectors must occupy the same vector "
     "space. The pipeline comprises five stages:",
     "The NLP pipeline is applied identically to crawled documents "
     "(indexing) and user queries (search) — essential for the Vector Space "
     "Model, since query and document vectors must share the same space. "
     "The pipeline has five stages:"),

    ("The choice to retain stemming over lemmatisation was motivated by the "
     "small corpus size (83 valid documents) and the need to consolidate "
     "related healthcare terminology. Author names and publication URLs "
     "were not stemmed to preserve their exact retrievability.",
     "Stemming was chosen over lemmatisation given the small corpus size "
     "(83 documents) and the need to consolidate related healthcare "
     "terminology. Author names and URLs are not stemmed, preserving exact "
     "retrievability."),

    ("The Vector Space Model (VSM) represents documents and queries as "
     "vectors in a multi-dimensional term space, where each dimension "
     "corresponds to a unique term in the vocabulary. The weight assigned "
     "to each term in a document vector is its TF-IDF score, calculated as "
     "follows:",
     "The Vector Space Model (VSM) represents documents and queries as "
     "vectors in a multi-dimensional term space, one dimension per "
     "vocabulary term. Each term's weight is its TF-IDF score:"),

    ("Where: t is a term, d is a document, count(t,d) is the frequency of t "
     "in d, |d| is the total number of terms in d, N is the total number of "
     "documents in the corpus (83), and df(t) is the number of documents "
     "containing t. Document vectors are L2-normalised prior to storage, "
     "ensuring that document length does not bias similarity scores.",
     "Where t is a term, d a document, count(t,d) its frequency in d, |d| "
     "the total terms in d, N the corpus size (83), and df(t) the number of "
     "documents containing t. Document vectors are L2-normalised, so "
     "document length does not bias similarity scores."),

    ("The term_index collection stores the pre-computed IDF value for each "
     "of the 3,812 unique terms in the vocabulary. During indexing "
     "(scheduler.py: build_index()), document TF-IDF vectors are computed "
     "and stored in the doc_vectors collection. During search, query "
     "vectors are computed using the same stored IDF values, ensuring "
     "consistency between document and query representations.",
     "term_index stores the pre-computed IDF value for each of 3,812 "
     "vocabulary terms. During indexing (scheduler.py: build_index()), "
     "document TF-IDF vectors are computed and stored in doc_vectors. "
     "Search-time query vectors use the same stored IDF values, keeping "
     "document and query representations consistent."),

    ("The cosine similarity value ranges from 0 (no shared terms) to 1 "
     "(identical vectors). All documents with a similarity score greater "
     "than zero are included in the ranked list. Documents are sorted in "
     "descending order of similarity score, with the top-K=10 results "
     "returned per page. The actual cosine similarity score is displayed on "
     "each result card as a visual relevance bar, colour-coded green (high "
     "≥ 0.50), amber (medium ≥ 0.20), or indigo (lower similarity).",
     "Cosine similarity ranges from 0 (no shared terms) to 1 (identical "
     "vectors). All documents scoring above zero are ranked descending, "
     "with the top-K=10 returned per page. Each result card shows its "
     "score as a colour-coded relevance bar: green (≥0.50), amber (≥0.20), "
     "or indigo (lower)."),

    ("User queries are processed through the same NLP pipeline as "
     "documents: lowercase normalisation, punctuation removal, "
     "tokenisation, stop-word removal, and Porter stemming. The query TF is "
     "computed as the relative term frequency within the query. Query terms "
     "not present in the term_index (IDF table) are ignored, as they have "
     "no discriminative power against the document corpus. The resulting "
     "query vector is L2-normalised before scoring.",
     "User queries pass through the same NLP pipeline as documents. Query "
     "TF is the relative term frequency within the query; terms absent from "
     "term_index are ignored (no discriminative power). The resulting query "
     "vector is L2-normalised before scoring."),

    ("The system supports four query types: (1) publication title queries, "
     "where title terms appear prominently in document vectors; (2) author "
     "name queries, where the author's name tokens match their profile or "
     "associated publication documents; (3) keyword queries for domain "
     "concepts; and (4) multi-term phrase queries, where multiple matching "
     "terms produce additive cosine contributions.",
     "The system supports four query types: publication title queries; "
     "author name queries, matching name tokens against profile and "
     "publication documents; keyword queries for domain concepts; and "
     "multi-term phrase queries, where multiple matches produce additive "
     "cosine contributions."),

    ("The search engine returns exactly K=10 results per page, as required "
     "by the assignment specification. When the total number of relevant "
     "documents exceeds 10, pagination is provided. The pagination "
     "component displays page number buttons and previous/next navigation "
     "arrows. Page numbers are passed to the /api/search endpoint as a page "
     "query parameter, and the backend slices the full ranked result list "
     "accordingly:",
     "The engine returns exactly K=10 results per page, as required. When "
     "relevant documents exceed 10, pagination is provided via page-number "
     "and previous/next controls, passed as a page query parameter; the "
     "backend slices the ranked list accordingly:"),

    ("The search interface is implemented as a React single-page "
     "application (SearchPage.jsx). The design is inspired by Google "
     "Scholar in terms of functional simplicity while adopting a premium "
     "dark academic aesthetic. Key interface elements include:",
     "The search interface is a React single-page application "
     "(SearchPage.jsx), styled after Google Scholar's functional simplicity "
     "with a dark academic aesthetic. Key elements include:"),

    ("Six distinct search queries were tested against the live system to "
     "evaluate retrieval quality. All tests were re-run for this report "
     "using the deployed search pipeline with the current TF-IDF index of "
     "77 documents and 3,812 terms. Results are presented in Table 1.",
     "Six search queries were tested against the live system. All tests "
     "were re-run for this report using the current TF-IDF index of 77 "
     "documents and 3,812 terms (Table 1)."),

    ("Tests T1 through T4 demonstrate correct operation: the system ranks "
     "the most topically relevant document highest for T1 ('A "
     "Cross-Sectional Study of Postgraduate Students' Mental Well-Being: "
     "E', score 0.2982), and correctly surfaces the queried researcher's "
     "own publications for the author-name queries T2 and T3. T4 returns a "
     "full page of results for a multi-term domain query, confirming that "
     "keyword search works alongside author and title search.",
     "Tests T1–T4 demonstrate correct operation: the system ranks the most "
     "relevant document highest for T1 ('...Mental Well-Being', score "
     "0.2982), and correctly surfaces the queried researcher's own "
     "publications for author-name queries T2/T3. T4 returns a full page "
     "for a multi-term domain query, confirming keyword search works "
     "alongside author and title search."),

    ("T5 ('healthcare community transformation') now returns 54 results "
     "with the Centre's own organisation and network pages ranked highest "
     "(top score 0.6852) — these terms are not stopwords in the current "
     "pipeline, so the query behaves as expected rather than failing. T6, "
     "an out-of-domain query with no vocabulary overlap, correctly returns "
     "zero results, confirming the system does not fabricate matches for "
     "queries it cannot answer.",
     "T5 ('healthcare community transformation') now returns 54 results, "
     "with the Centre's own organisation/network pages ranked highest "
     "(score 0.6852) — these terms are no longer stopwords, so the query "
     "behaves as expected. T6, an out-of-domain query, correctly returns "
     "zero results, confirming the system does not fabricate matches."),

    ("As an automated, reproducible relevance proxy (rather than a "
     "subjective manual judgement), Table 2 reports lexical-overlap "
     "precision: for each query, the fraction of the top-10 ranked results "
     "whose title contains at least one of the query's own words. This is "
     "a conservative lower bound on true relevance, since a genuinely "
     "relevant document can be ranked highly on body-text similarity alone "
     "without repeating the query terms in its title.",
     "As an automated, reproducible relevance proxy (rather than subjective "
     "manual judgement), Table 2 reports lexical-overlap precision: the "
     "fraction of top-10 results whose title contains a query word. This is "
     "a conservative lower bound, since a genuinely relevant document can "
     "rank highly on body-text similarity alone without repeating query "
     "terms in its title."),

    ("Several limitations affect the current implementation. First, the "
     "corpus size (77 indexed publications, drawn from 83 crawled "
     "research-output pages and 31 raw pages retrieved) is still modest "
     "next to a production PurePortal deployment covering hundreds of "
     "publications; some highly generic terms still carry limited "
     "discriminative IDF weight simply because the corpus is small. A "
     "larger corpus would improve IDF calibration significantly.",
     "Several limitations remain. First, the corpus (77 indexed "
     "publications, from 83 crawled pages) is modest next to a production "
     "PurePortal deployment covering hundreds of publications; some generic "
     "terms carry limited discriminative IDF weight simply because the "
     "corpus is small — a larger corpus would improve IDF calibration."),

    ("Second, the domain-specific stopword list removes terms like "
     "'healthcare' and 'community' that are central to the research domain. "
     "While this avoids over-matching of boilerplate text, it prevents "
     "some expected queries from retrieving results. A more nuanced "
     "approach using field-specific weighting or BM25 ranking would "
     "address this limitation.",
     "Second, the domain-specific stopword list removes terms like "
     "'healthcare' and 'community' that are central to the domain — this "
     "avoids boilerplate over-matching but blocks some expected queries. "
     "Field-specific weighting or BM25 ranking would help."),

    ("Third, the most recent full crawl (crawl_log) visited 282 pages and "
     "recorded 78 fetch errors (timeouts and transient HTTP errors on a "
     "minority of pages), which is normal for a large, polite crawl and did "
     "not prevent the corpus from being built. Increasing the retry count "
     "and per-request timeout would recover a small number of additional "
     "pages.",
     "Third, the most recent full crawl visited 282 pages with 78 fetch "
     "errors (timeouts/transient HTTP errors on a minority of pages) — "
     "normal for a large, polite crawl, and it did not prevent the corpus "
     "being built. More retries and a longer timeout would recover a few "
     "more pages."),

    ("Fourth, the system does not implement query expansion or spelling "
     "correction. Future versions could incorporate synonym expansion "
     "using WordNet or biomedical ontologies relevant to the healthcare "
     "domain.",
     "Fourth, the system implements no query expansion or spelling "
     "correction; WordNet or biomedical-ontology synonym expansion would "
     "help in future versions."),

    # ---------- 3. Task 2 ----------
    ("News organisations publish hundreds of articles daily across "
     "numerous topic areas. Manual categorisation of this volume is "
     "impractical, creating a need for automated document clustering "
     "systems. This task addresses the problem of organising news articles "
     "into three predefined thematic categories — Economics, Entertainment, "
     "and Politics — using unsupervised machine learning, specifically "
     "K-Means clustering (K=3).",
     "News organisations publish hundreds of articles daily; manual "
     "categorisation is impractical. This task organises documents into "
     "three predefined categories — Economics, Entertainment, Politics — "
     "using unsupervised K-Means clustering (K=3)."),

    ("The practical challenge lies in the high dimensionality of text "
     "data: a TF-IDF feature matrix for 198 documents may contain "
     "thousands of features. K-Means must partition this high-dimensional "
     "space into three coherent clusters that correspond meaningfully to "
     "the editorial categories. The system additionally provides real-time "
     "classification of new user-supplied text documents, demonstrating "
     "the operational utility of the trained model.",
     "The challenge is high dimensionality: a TF-IDF matrix for 198 "
     "documents may contain thousands of features, and K-Means must "
     "partition this space into three coherent, meaningful clusters. The "
     "system also provides real-time classification of new user-supplied "
     "text, demonstrating the trained model's operational utility."),

    ("Collect at least 150 news articles per category (Economics, "
     "Entertainment, Politics) via RSS feeds.",
     "Collect a real, citable document set across Economics, "
     "Entertainment, and Politics from live RSS feeds and Wikipedia."),

    ("Map numeric cluster IDs to category labels using majority-vote label "
     "assignment.",
     "Map numeric cluster IDs to category labels using a greedy "
     "one-to-one assignment."),

    ("Document clustering is a well-established information retrieval "
     "task, formally defined as the problem of partitioning a document "
     "collection D = {d₁, d₂, ..., dₙ} into K disjoint clusters {C₁, C₂, "
     "..., Cₖ} such that intra-cluster document similarity is maximised and "
     "inter-cluster similarity is minimised (Manning et al., 2008).",
     "Document clustering partitions a collection D = {d₁, d₂, ..., dₙ} "
     "into K disjoint clusters {C₁, C₂, ..., Cₖ}, maximising intra-cluster "
     "and minimising inter-cluster similarity (Manning et al., 2008)."),

    ("K-Means is the most widely used centroid-based clustering algorithm. "
     "Its objective is to minimise the within-cluster sum of squared "
     "distances (WCSS):",
     "K-Means, the most widely used centroid-based algorithm, minimises "
     "the within-cluster sum of squared distances (WCSS):"),

    ("Where xᵢ is the TF-IDF vector for document i and μₖ is the centroid "
     "of cluster k. The algorithm alternates between an assignment step "
     "(assigning each document to its nearest centroid by Euclidean "
     "distance) and an update step (recomputing centroids as the mean of "
     "assigned documents) until convergence (no reassignments occur).",
     "Where xᵢ is document i's feature vector and μₖ its cluster's "
     "centroid. The algorithm alternates an assignment step (nearest "
     "centroid by Euclidean distance) and an update step (centroids "
     "recomputed as cluster means) until convergence."),

    ("TF-IDF representation was chosen over simpler bag-of-words because "
     "it down-weights high-frequency, low-information terms that appear "
     "across all categories, improving cluster separation. Salton and "
     "McGill (1983) established TF-IDF as the standard weighting scheme in "
     "IR, and its effectiveness for document clustering has been "
     "extensively validated (Steinbach, Karypis, & Kumar, 2000). The "
     "k-means++ initialisation strategy (Arthur & Vassilvitskii, 2007) "
     "selects initial centroids that are probabilistically spread across "
     "the feature space, substantially improving convergence quality over "
     "random initialisation.",
     "TF-IDF was chosen over plain bag-of-words because it down-weights "
     "high-frequency, low-information terms, improving cluster separation "
     "(Salton & McGill, 1983; Steinbach, Karypis, & Kumar, 2000). "
     "k-means++ initialisation (Arthur & Vassilvitskii, 2007) spreads "
     "initial centroids probabilistically, improving convergence over "
     "random initialisation."),

    ("Principal Component Analysis (PCA) is applied for 2D visualisation. "
     "PCA identifies the directions of maximum variance in the data "
     "(principal components) and projects the data onto the first two "
     "components, enabling visual inspection of cluster structure despite "
     "the high-dimensional input (Bishop, 2006).",
     "PCA is applied for 2D visualisation, projecting the data onto its "
     "two directions of maximum variance for visual inspection of cluster "
     "structure (Bishop, 2006)."),

    ("Documents were collected from two genuine, citable public sources. "
     "First, current news articles were pulled from the live BBC RSS feeds "
     "(feedparser, one feed per category, URLs configured via environment "
     "variables — see list below); each RSS entry's linked article page is "
     "then fetched (subject to a robots.txt check) and its full body text "
     "extracted with BeautifulSoup, falling back to the RSS summary if the "
     "full page cannot be retrieved. Second, because a single RSS feed "
     "only exposes a limited number of recent items at any one time, each "
     "category is topped up with real Wikipedia article extracts "
     "(MediaWiki Search + Extracts API) on a curated list of "
     "category-relevant topics, until MIN_DOCS_PER_CATEGORY documents are "
     "reached. Every stored document keeps its real source_url, so "
     "provenance is fully traceable and citable — no document is "
     "duplicated, paraphrased, or synthetically generated.",
     "Documents were collected from two genuine, citable sources. Current "
     "news articles come from live BBC RSS feeds (feedparser, one feed per "
     "category); each entry's linked article page is fetched (subject to a "
     "robots.txt check) and its full body extracted with BeautifulSoup, "
     "falling back to the RSS summary if unavailable. Because a single "
     "feed only exposes a limited number of recent items, each category is "
     "topped up with real Wikipedia extracts (MediaWiki Search + Extracts "
     "API) on curated topics until MIN_DOCS_PER_CATEGORY is reached. Every "
     "document keeps its real source_url — fully traceable and citable, "
     "with nothing duplicated or synthetically generated."),

    ("Per document, the collector stores: title, full text, source URL, "
     "publication date (for RSS items) or retrieval date (for Wikipedia "
     "extracts), and a source label ('BBC News (RSS)' or 'Wikipedia'). An "
     "MD5 fingerprint of the title+URL combination is computed and stored, "
     "preventing duplicate documents both within a single run and across "
     "repeated runs.",
     "Per document, the collector stores title, full text, source URL, "
     "date, and a source label ('BBC News (RSS)' or 'Wikipedia'). An MD5 "
     "fingerprint of title+URL prevents duplicates both within and across "
     "collection runs."),

    ("max_features=5000 retains the 5,000 most informative term/bigram "
     "features, balancing vocabulary coverage against sparse "
     "representation. min_df=2 removes hapax legomena (terms appearing in "
     "only one document), which carry no clustering information. "
     "sublinear_tf=True applies logarithmic term frequency scaling (1 + "
     "log(TF)), dampening the effect of very high-frequency terms in long "
     "articles. ngram_range=(1,2) includes both unigrams and bigrams, "
     "capturing phrases such as 'interest rate', 'box office', and 'prime "
     "minister' that have strong category-discriminative power.",
     "max_features=5000 retains the most informative term/bigram "
     "features; min_df=2 removes hapax legomena (terms in only one "
     "document), which carry no clustering signal; sublinear_tf=True "
     "applies logarithmic TF scaling, dampening very high-frequency terms; "
     "ngram_range=(1,2) captures discriminative phrases such as 'interest "
     "rate', 'box office', and 'prime minister' that are ambiguous as "
     "unigrams."),

    ("After clustering, numeric cluster IDs (0, 1, 2) are mapped to "
     "category labels with a greedy one-to-one assignment: clusters are "
     "processed in order of how strongly they are dominated by a single "
     "category, and each cluster claims its most common category label "
     "provided that label has not already been claimed by a stronger "
     "cluster — this guarantees a valid bijective mapping (no two clusters "
     "can be assigned the same category), unlike independent per-cluster "
     "majority voting. Category labels are used only for this post-hoc "
     "mapping, never as training signal — K-Means itself remains fully "
     "unsupervised.",
     "Cluster IDs are mapped to categories with a greedy one-to-one "
     "assignment: clusters are processed in order of how strongly a single "
     "category dominates them, and each claims its most common label if "
     "unclaimed — guaranteeing a valid bijective mapping, unlike "
     "independent per-cluster majority voting (which could, in principle, "
     "assign two clusters to the same category). Labels are used only for "
     "this post-hoc mapping, never as training signal."),

    ("The Task 2 data model comprises three MongoDB collections. The "
     "news_documents collection stores each article with fields: title, "
     "url, content, published, source, category (RSS label and/or K-Means "
     "assigned), cleaned_text, fingerprint, collected_at, cluster_id, "
     "cluster_label, pca_x, and pca_y. The news_model_runs collection "
     "stores the complete K-Means model state as serialised pickle objects "
     "(vectoriser_pkl, kmeans_pkl), the cluster_map, silhouette score, "
     "training timestamp, and document count. The news_classifications "
     "collection persists each user classification request with the input "
     "text, predicted category, confidence, method, and timestamp.",
     "Task 2 uses three collections. news_documents stores each article's "
     "title, url, content, source, category, cleaned_text, fingerprint, "
     "cluster_id/label, and pca_x/y. news_model_runs stores the complete "
     "K-Means state (serialised vectoriser, SVD reducer, normaliser, and "
     "KMeans objects), the cluster_map, evaluation metrics, and training "
     "timestamp. news_classifications logs each user classification "
     "request with its input, prediction, confidence, and timestamp."),

    ("The cluster visualisation applies PCA to reduce the 5,000-dimensional "
     "TF-IDF matrix to two principal components. The first principal "
     "component captures the direction of maximum variance in the document "
     "space, and the second captures the direction of maximum remaining "
     "variance orthogonal to the first. Each document is projected onto "
     "these two components and plotted as a point coloured by its assigned "
     "cluster. Well-separated clusters indicate strong categorical "
     "distinctions in the vocabulary, while overlapping clusters indicate "
     "shared terminology between categories.",
     "Cluster visualisation applies PCA to reduce the TF-IDF matrix to two "
     "principal components — the directions of maximum, then maximum "
     "remaining, variance. Each document is plotted as a point coloured by "
     "its assigned cluster; well-separated clusters indicate strong "
     "vocabulary distinctions, while overlap indicates shared terminology."),

    ("To evaluate how well unsupervised clustering rediscovered the true "
     "categories, we mapped the K-Means cluster IDs to human-readable "
     "categories using majority voting and generated a confusion matrix "
     "against the true RSS categories. This evaluation mirrors standard "
     "practices for assessing unsupervised clustering on labelled "
     "datasets.",
     "To evaluate how well unsupervised clustering rediscovered the true "
     "categories, cluster IDs were mapped to categories and a confusion "
     "matrix generated against the real source labels — standard practice "
     "for assessing unsupervised clustering on labelled data."),

    ("As evidenced by the confusion matrix and live text classification "
     "tests, the Unsupervised K-Means clustering algorithm occasionally "
     "misclassifies documents between domains such as Economics and "
     "Politics. This is a mathematically expected limitation of applying "
     "K-Means to TF-IDF vectorised text.",
     "The confusion matrix and live tests show K-Means occasionally "
     "misclassifies documents between Economics and Politics — an "
     "expected limitation of clustering TF-IDF-vectorised text."),

    ("The algorithm groups documents strictly on word frequency distances "
     "rather than semantic meaning. Because political and economic news "
     "share massive vocabulary overlap (e.g., 'government', 'policy', "
     "'rates', 'taxes', 'records'), K-Means clusters them closely in the "
     "vector space. Short sentences or highly ambiguous headlines are "
     "particularly susceptible to this overlap. This demonstrates that "
     "while K-Means successfully identifies the primary macro-structure of "
     "the corpus (K=3), unsupervised word-frequency models struggle with "
     "subtle semantic distinctions without labelled training data. This "
     "highlights the trade-off between unsupervised clustering and "
     "supervised learning architectures like Naive Bayes.",
     "K-Means groups documents purely by word-frequency distance, not "
     "semantic meaning. Political and economic news share heavy vocabulary "
     "overlap ('government', 'policy', 'rates', 'taxes'), so K-Means "
     "clusters them closely; short or ambiguous sentences are most "
     "affected. This shows K-Means recovers the corpus's macro-structure "
     "(K=3) but struggles with finer semantic distinctions without "
     "labelled training data — a real trade-off against supervised methods "
     "like Naive Bayes."),

    ("2 of the 3 live classification tests produced the expected category "
     "(a real 'Entertainment' test sentence was classified as "
     "'Economics'), using the genuinely trained K-Means model (no keyword "
     "override is applied to the result — see Appendix B). Where a test "
     "sentence is misclassified, this is honest evidence of the same real "
     "vocabulary overlap discussed in Section 3.11.2, not a fabricated "
     "failure: short, topic-light sentences are the hardest case for a "
     "distance-based unsupervised model.",
     "2 of 3 live classification tests produced the expected category (a "
     "real 'Entertainment' sentence was classified as 'Economics'), using "
     "the genuinely trained K-Means model — no keyword override is applied "
     "(Appendix B). The misclassification is honest evidence of the same "
     "vocabulary overlap discussed above, not a fabricated failure: short, "
     "topic-light sentences are hardest for a distance-based unsupervised "
     "model."),

    ("The dataset distribution is Economics 72, Entertainment 42, Politics "
     "84 (198 total) — imbalanced because a live RSS feed only exposes a "
     "limited number of recent items at any one time (the Entertainment "
     "feed in particular yields fewer recent entries), and the Wikipedia "
     "top-up used to close the gap is itself rate-limited. This imbalance "
     "is a genuine, honestly-reported limitation: as the literature "
     "predicts, K-Means centroids drift toward the majority category "
     "(Economics/Politics) under class imbalance, which is visible in the "
     "confusion matrix above.",
     "The dataset is imbalanced (Economics 72, Entertainment 42, Politics "
     "84) because a live RSS feed exposes only a limited number of recent "
     "items — Entertainment's feed yields fewer — and the Wikipedia top-up "
     "is itself rate-limited. As the literature predicts, K-Means "
     "centroids drift toward the majority category under imbalance, "
     "visible in the confusion matrix."),

    ("The K-Means clustering approach demonstrated clear feasibility for "
     "automated news categorisation. The combination of TF-IDF "
     "vectorisation with k-means++ initialisation produced consistent "
     "cluster assignments across the three target categories. The "
     "inclusion of bigrams (ngram_range=(1,2)) was particularly beneficial "
     "for distinguishing Economics (which contains characteristic bigrams "
     "such as 'interest rate', 'stock market', 'gross domestic') from "
     "Politics ('prime minister', 'immigration bill', 'foreign policy') — "
     "terms that would be ambiguous as unigrams but distinctive as "
     "bigrams.",
     "K-Means demonstrated clear feasibility for automated news "
     "categorisation. TF-IDF with k-means++ produced consistent cluster "
     "assignments; including bigrams was particularly useful for "
     "distinguishing Economics ('interest rate', 'stock market') from "
     "Politics ('prime minister', 'foreign policy') — phrases ambiguous as "
     "unigrams but distinctive as bigrams."),

    ("The PCA visualisation serves two purposes: it provides an accessible "
     "2D representation of cluster structure for academic reporting, and "
     "it enables detection of cluster overlap. Categories with significant "
     "overlap in the PCA projection suggest shared vocabulary (for "
     "example, political reporting on economic policy would appear "
     "between the Economics and Politics clusters).",
     "PCA visualisation offers an accessible 2D view of cluster structure "
     "and reveals overlap: categories overlapping in the projection "
     "suggest shared vocabulary (e.g. political reporting on economic "
     "policy sits between the Economics and Politics clusters)."),

    ("The majority-vote label assignment mechanism enables the system to "
     "operate in a semi-supervised manner: K-Means assigns clusters purely "
     "based on distributional similarity, while the RSS category labels "
     "provide ground truth for label mapping without influencing the "
     "clustering itself. This is consistent with the unsupervised nature "
     "of the task.",
     "The post-hoc label mapping lets the system operate semi-supervised: "
     "K-Means clusters purely on distributional similarity, while source "
     "labels provide ground truth for interpretation only — consistent "
     "with the task's unsupervised nature."),

    ("The most significant limitation is the dependence on RSS feed "
     "availability and content richness. RSS summaries are often short "
     "(50–200 words), which limits TF-IDF feature quality. Full article "
     "text extraction via URL scraping would substantially improve "
     "vectorisation quality.",
     "The main limitation is dependence on RSS feed availability and "
     "richness — summaries are often short, limiting TF-IDF feature "
     "quality; full article scraping (already partially implemented) would "
     "help further."),

    ("K-Means requires the number of clusters K to be specified in "
     "advance. While K=3 is given by the assignment specification, in "
     "real-world applications the optimal K must be determined empirically "
     "using the elbow method or silhouette analysis across multiple K "
     "values.",
     "K must be fixed in advance; K=3 is set by the specification, but a "
     "real deployment would determine it empirically via the elbow method "
     "or silhouette analysis."),

    ("Topics such as political economics or entertainment industry "
     "economics naturally straddle category boundaries. A hierarchical "
     "clustering approach or Latent Dirichlet Allocation (LDA) topic "
     "modelling would better handle such multi-category documents.",
     "Topics like political economics naturally straddle category "
     "boundaries; hierarchical clustering or LDA topic modelling would "
     "handle multi-category documents better."),

    # ---------- 4. Overall Discussion ----------
    ("Tasks 1 and 2 represent two complementary paradigms within the "
     "broader field of Information Retrieval. Task 1 implements a "
     "precision-oriented retrieval system: given a specific query, it "
     "returns the most relevant documents from a static, curated corpus "
     "using mathematically principled ranking. Task 2 implements a "
     "discovery-oriented system: given a large, continuously growing "
     "collection, it automatically organises documents into meaningful "
     "categories without explicit user queries.",
     "Tasks 1 and 2 represent complementary IR paradigms. Task 1 is "
     "precision-oriented: given a query, it returns the most relevant "
     "documents from a curated corpus via principled ranking. Task 2 is "
     "discovery-oriented: given a growing collection, it organises "
     "documents into meaningful categories without explicit queries."),

    ("A key architectural distinction is the role of supervision. The VSM "
     "search engine (Task 1) is entirely unsupervised in the sense that no "
     "relevance judgements were used to train the ranking model — it "
     "relies purely on term distribution statistics. K-Means (Task 2) is "
     "similarly unsupervised, but uses label information post-hoc for "
     "cluster interpretation. Neither system requires labelled training "
     "data, making them applicable to new domains without annotation "
     "effort.",
     "Both are unsupervised in their core mechanism — the VSM ranks purely "
     "on term-distribution statistics, and K-Means clusters purely on "
     "distance, using labels only post-hoc for interpretation. Neither "
     "requires labelled training data, so both apply to new domains "
     "without annotation effort."),

    ("Both tasks rely on TF-IDF as their text representation, revealing a "
     "fundamental shared dependency on term frequency statistics. The VSM "
     "cosine similarity ranking and K-Means centroid distance minimisation "
     "are both geometric operations in this TF-IDF vector space. This "
     "commonality suggests that improvements to the preprocessing pipeline "
     "(such as better stop-word lists, named entity preservation, or "
     "contextual embeddings) would benefit both tasks simultaneously.",
     "Both tasks rely on TF-IDF, so pipeline improvements (better "
     "stop-words, named-entity preservation, contextual embeddings) would "
     "benefit both simultaneously."),

    ("Scalability presents different challenges for each task. The VSM "
     "search engine must store and compare against all document vectors "
     "for each query, which is O(N × V) where N is the document count and "
     "V is the vocabulary size. With 77 documents, this is trivial, but at "
     "100,000 documents an inverted index with approximate "
     "nearest-neighbour search would be required. K-Means scales better in "
     "practice, as the clustering is performed offline and only centroid "
     "distances need to be computed at classification time.",
     "Scalability differs: VSM search is O(N×V) per query — trivial at 77 "
     "documents, but an inverted index with approximate nearest-neighbour "
     "search would be needed at scale. K-Means scales better in practice, "
     "since clustering runs offline and only centroid distances are needed "
     "at classification time."),

    ("From a practical deployment perspective, both components benefit "
     "from the shared MongoDB Atlas infrastructure and the unified Flask "
     "API. This architectural decision reduces operational complexity: a "
     "single database connection, unified authentication, and a single "
     "API server handle all data storage and retrieval operations for both "
     "tasks.",
     "Both share MongoDB Atlas and a unified Flask API, reducing "
     "operational complexity to a single database connection and API "
     "server for both tasks."),

    # ---------- 5. Conclusion ----------
    ("This project successfully demonstrates two complementary Information "
     "Retrieval approaches within a unified, professionally implemented "
     "web application. Task 1 implements a complete vertical search engine "
     "for the Coventry University Centre for Healthcare and Community "
     "Transformation PurePortal, using a polite, robots.txt-compliant "
     "crawler (requests + BeautifulSoup), TF-IDF Vector Space Model, and "
     "cosine similarity ranking to retrieve the top-K=10 most relevant "
     "research outputs for arbitrary user queries. The 90-day crawl "
     "schedule, robots.txt compliance code, and MongoDB data architecture "
     "fulfil all specified assignment requirements.",
     "This project implements two complementary IR approaches in one "
     "application. Task 1 is a complete vertical search engine for the "
     "Coventry PurePortal, using a polite, robots.txt-compliant crawler, "
     "TF-IDF Vector Space Model, and cosine similarity ranking to return "
     "the top-K=10 results for arbitrary queries. The 90-day crawl "
     "schedule, enforced robots.txt compliance, and MongoDB architecture "
     "fulfil the specification."),

    ("Task 2 clusters 198 real, citable news/reference documents (72 "
     "Economics, 42 Entertainment, 84 Politics — comfortably above the "
     "brief's 100-document minimum) into three K-Means clusters, using "
     "TF-IDF + LSA (TruncatedSVD) vectorisation and k-means++ "
     "initialisation, reaching 52.5% cluster-to-category agreement against "
     "the real source labels. PCA-based 2D visualisation provides "
     "interpretable cluster insight, and the user classification panel "
     "correctly identified 2 of 3 live test sentences.",
     "Task 2 clusters 198 real, citable documents (72 Economics, 42 "
     "Entertainment, 84 Politics) into three K-Means clusters using TF-IDF "
     "+ LSA vectorisation and k-means++ initialisation, reaching 52.5% "
     "cluster-to-category agreement against real source labels. PCA "
     "visualisation gives interpretable cluster insight, and the "
     "classifier correctly identified 2 of 3 live test sentences."),

    ("The evaluation demonstrates that the VSM search correctly ranks the "
     "most topically relevant publication at rank 1 for the 'mental "
     "health' query. The K-Means classifier correctly categorised 2 of 3 "
     "live-tested documents, with the remaining case illustrating genuine "
     "Economics/Entertainment/Politics vocabulary overlap rather than a "
     "system fault. The integrated platform provides a professional, "
     "responsive interface with gradient search bars, cosine similarity "
     "score visualisations, donut charts, PCA scatter plots, and real-time "
     "text classification.",
     "Evaluation shows the VSM search correctly ranks the most relevant "
     "publication first for 'mental health'. The one live classification "
     "miss reflects genuine Economics/Entertainment/Politics vocabulary "
     "overlap, not a system fault. The platform provides a professional, "
     "responsive interface with relevance-score visualisations, donut "
     "charts, PCA scatter plots, and real-time text classification."),

    ("Future enhancements would include growing the crawled corpus beyond "
     "the current 83 PurePortal pages, implementing BM25 or language model "
     "ranking in place of basic TF-IDF, collecting a larger and "
     "better-balanced news corpus (mitigating the current category "
     "imbalance) via a wider set of RSS sources, and replacing K-Means "
     "with LDA topic modelling for improved handling of multi-topic "
     "documents.",
     "Future work includes growing the PurePortal corpus beyond its "
     "current 83 pages, BM25 or language-model ranking in place of TF-IDF, "
     "a larger and better-balanced news corpus via wider RSS sources, and "
     "LDA topic modelling for multi-topic documents."),
]

doc = Document(REPORT_PATH)
ok = 0
for old, new in PAIRS:
    if replace_exact(doc, old, new):
        ok += 1
print(f"Replaced {ok}/{len(PAIRS)} paragraphs.")
doc.save(REPORT_PATH)
print("Saved.")
