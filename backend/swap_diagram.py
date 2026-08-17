import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")
IMG_PATH = os.path.join(os.path.dirname(__file__), "static", "conceptual_diagram.png")


def paragraph_has_image(p):
    return bool(p._element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"))


def clear_paragraph(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


doc = Document(REPORT_PATH)
target = None
for p in doc.paragraphs:
    if paragraph_has_image(p):
        target = p
        break  # the diagram is the first image in the document

if target is not None:
    clear_paragraph(target)
    run = target.add_run()
    run.add_picture(IMG_PATH, width=Inches(6.3))
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print("Replaced conceptual diagram (first image in doc).")
else:
    print("No image found.")

doc.save(REPORT_PATH)
print("Saved.")
