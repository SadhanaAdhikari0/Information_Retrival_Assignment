"""fix_report_pass9.py — last three stale-number/claim stragglers."""
import os
from docx import Document
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), ".env"))
from pymongo import MongoClient

db = MongoClient(os.environ["MONGODB_URI"])["vertical_search_engine"]
REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")

n_research_outputs = db.research_outputs.count_documents({})
n_profiles = db.researcher_profiles.count_documents({})
n_doc_vectors = db.doc_vectors.count_documents({})
n_term_index = db.term_index.count_documents({})


def set_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def replace_in_all_text(doc, replacements):
    def process(paragraph):
        full = "".join(r.text for r in paragraph.runs)
        if not full:
            return
        new_full = full
        for old, new in replacements:
            new_full = new_full.replace(old, new)
        if new_full != full:
            set_paragraph_text(paragraph, new_full)
    for p in doc.paragraphs:
        process(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process(p)


FIXES = [
    ("The Task 1 architecture follows a layered pipeline pattern. The data "
     "layer comprises MongoDB Atlas collections: doc_vectors (17 TF-IDF "
     "document vectors), term_index (2,174 indexed terms with IDF values), "
     "research_outputs (7 publications), researcher_profiles (10 "
     "profiles), and crawl_log (crawl history).",
     "The Task 1 architecture follows a layered pipeline pattern. The data "
     f"layer comprises MongoDB Atlas collections: doc_vectors "
     f"({n_doc_vectors} TF-IDF document vectors), term_index "
     f"({n_term_index:,} indexed terms with IDF values), research_outputs "
     f"({n_research_outputs} publications), researcher_profiles "
     f"({n_profiles} profiles), and crawl_log (crawl history)."),

    ("With 17 documents, this is trivial, but at 100,000 documents an "
     "inverted index with approximate nearest-neighbour search would be "
     "required.",
     f"With {n_doc_vectors} documents, this is trivial, but at 100,000 "
     "documents an inverted index with approximate nearest-neighbour "
     "search would be required."),

    ("Task 1 implements a complete vertical search engine for the Coventry "
     "University Centre for Healthcare and Community Transformation "
     "PurePortal, using a Selenium-based crawler, TF-IDF Vector Space "
     "Model, and cosine similarity ranking",
     "Task 1 implements a complete vertical search engine for the Coventry "
     "University Centre for Healthcare and Community Transformation "
     "PurePortal, using a polite, robots.txt-compliant crawler "
     "(requests + BeautifulSoup), TF-IDF Vector Space Model, and cosine "
     "similarity ranking"),
]
doc = Document(REPORT_PATH)
replace_in_all_text(doc, FIXES)
doc.save(REPORT_PATH)
print("Saved pass 9.")
