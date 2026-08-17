"""
swap_code_images.py — Replaces the embedded code-evidence screenshots in
the report with the freshly regenerated ones (screenshots/*.png), which
were rendered directly from the real, current source code.

Finds each image by locating the paragraph immediately preceding it (a
known caption sentence), clears that picture paragraph, and inserts the
new PNG in its place.
"""
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")
SCREENSHOTS = os.path.join(os.path.dirname(__file__), "..", "screenshots")

# (anchor text of the paragraph BEFORE the image, screenshot filename)
REPLACEMENTS = [
    ("The following snippet demonstrates the Web Crawler data acquisition loop.",
     "crawler_code.png"),
    ("The following snippet demonstrates the automated scheduler running the "
     "crawler exactly every 3 months (90 days).", "scheduler_code.png"),
    ("The following snippet demonstrates the TF-IDF vectorisation and Cosine "
     "Similarity ranking logic for the Vertical Search Engine.",
     "search_engine_code.png"),
    ("The following snippet demonstrates scraping live news from BBC and The "
     "Guardian to build the dataset.", "task2_scraper_code.png"),
    ("The following snippet demonstrates the K-Means clustering training and "
     "evaluation pipeline.", "kmeans_training.png"),
    ("The following snippet demonstrates the greedy 1-to-1 cluster-to-category "
     "mapping logic.", "greedy_mapping.png"),
]


def find_paragraph_index(doc, text):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == text:
            return i
    return None


def paragraph_has_image(p):
    return bool(p._element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"))


def clear_paragraph(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


doc = Document(REPORT_PATH)

for anchor_text, filename in REPLACEMENTS:
    idx = find_paragraph_index(doc, anchor_text)
    if idx is None:
        print(f"  [WARN] anchor not found: {anchor_text[:60]}")
        continue

    # the image should be in one of the next couple of paragraphs
    img_para = None
    for j in range(idx + 1, min(idx + 4, len(doc.paragraphs))):
        if paragraph_has_image(doc.paragraphs[j]):
            img_para = doc.paragraphs[j]
            break

    img_path = os.path.join(SCREENSHOTS, filename)
    if not os.path.exists(img_path):
        print(f"  [WARN] screenshot missing: {img_path}")
        continue

    if img_para is not None:
        clear_paragraph(img_para)
        run = img_para.add_run()
        run.add_picture(img_path, width=Inches(6.0))
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f"  Replaced image after: {anchor_text[:50]}...")
    else:
        # No existing image paragraph — insert one right after the anchor.
        p = doc.paragraphs[idx]
        new_p = p.insert_paragraph_before("")  # placeholder position marker
        # Move new_p to just after p instead of before (python-docx has no
        # direct "insert after"), by moving its XML element.
        p._p.addnext(new_p._p)
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new_p.add_run()
        run.add_picture(img_path, width=Inches(6.0))
        print(f"  Inserted new image after: {anchor_text[:50]}...")

doc.save(REPORT_PATH)
print("Saved.")
