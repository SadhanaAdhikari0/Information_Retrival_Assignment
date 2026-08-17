import os
from docx import Document

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")


def set_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def replace_exact(doc, old, new, label=""):
    for p in doc.paragraphs:
        if p.text.strip() == old.strip():
            set_paragraph_text(p, new)
            return True
    print(f"  [WARN] not found: {label or old[:60]}")
    return False


def delete_exact(doc, text, label=""):
    for p in doc.paragraphs:
        if p.text.strip() == text.strip():
            p._element.getparent().remove(p._element)
            return True
    print(f"  [WARN] delete target not found: {label or text[:60]}")
    return False


PAIRS = [
    ("Ethical crawling is a core IR requirement. The crawler checks every "
     "URL against PurePortal's robots.txt via Python's urllib.robotparser "
     "before fetching. Building this surfaced a genuine bug: "
     "RobotFileParser.read() fetches robots.txt via bare urllib.request "
     "with no custom headers, and PurePortal's edge protection returns "
     "HTTP 403 to urllib's default User-Agent (confirmed independently — "
     "the identical URL returns 200 via requests). RobotFileParser "
     "silently treats a 403 as \"disallow everything\", so the first "
     "version of this check blocked the crawler from fetching even its "
     "own seed URL. The fix — fetch the robots.txt text via requests with "
     "the crawler's own User-Agent, then hand it to RobotFileParser.parse() "
     "instead of .read() — resolved it. PurePortal's robots.txt (checked "
     "2026-08-15) permits every path this crawler visits and specifies "
     "Crawl-Delay: 5, matching the crawler's delay.",
     "The crawler checks every URL against PurePortal's robots.txt via "
     "urllib.robotparser before fetching. This surfaced a genuine bug: "
     "RobotFileParser.read() fetches robots.txt via bare urllib.request "
     "with no custom headers, and PurePortal's edge protection 403s "
     "urllib's default User-Agent (confirmed: the identical URL returns "
     "200 via requests). RobotFileParser silently treats a 403 as "
     "\"disallow everything\", so the first version blocked the crawler "
     "from fetching even its own seed URL. Fix: fetch the text via "
     "requests with the crawler's own User-Agent, then parse() it instead "
     "of .read(). PurePortal's robots.txt (checked 2026-08-15) permits "
     "every path this crawler visits, with Crawl-Delay: 5 matching the "
     "crawler's own delay."),

    ("Documents were collected from two genuine, citable sources. Current "
     "news articles come from live BBC RSS feeds (feedparser, one feed per "
     "category); each entry's linked article page is fetched (subject to "
     "a robots.txt check) and its full body extracted with BeautifulSoup, "
     "falling back to the RSS summary if unavailable. Because a single "
     "feed only exposes a limited number of recent items, each category "
     "is topped up with real Wikipedia extracts (MediaWiki Search + "
     "Extracts API) on curated topics until MIN_DOCS_PER_CATEGORY is "
     "reached. Every document keeps its real source_url — fully traceable "
     "and citable, with nothing duplicated or synthetically generated.",
     "Documents come from two genuine, citable sources: live BBC RSS "
     "feeds (feedparser; each entry's article page is fetched, subject to "
     "a robots.txt check, and its full body extracted, falling back to "
     "the RSS summary if unavailable), topped up per category with real "
     "Wikipedia extracts (MediaWiki Search + Extracts API on curated "
     "topics) since a single feed only exposes limited recent items. "
     "Every document keeps its real source_url — nothing is duplicated or "
     "synthetically generated."),

    ("The Task 1 architecture is a layered pipeline. The data layer "
     "comprises MongoDB Atlas collections: doc_vectors (77 TF-IDF "
     "vectors), term_index (3,812 terms with IDF values), research_outputs "
     "(83 publications), researcher_profiles (122 profiles), and "
     "crawl_log. The crawler layer uses requests + BeautifulSoup for HTML "
     "parsing and metadata extraction. The processing layer applies the "
     "NLP pipeline and builds TF-IDF vectors. The API layer is a Flask "
     "REST API (/api/search, /api/suggestions, /api/crawl-status); the "
     "presentation layer is a React single-page application.",
     "Task 1 is a layered pipeline: MongoDB Atlas (doc_vectors, "
     "term_index, research_outputs, researcher_profiles, crawl_log) → a "
     "requests+BeautifulSoup crawler → an NLP/TF-IDF processing layer → a "
     "Flask REST API (/api/search, /api/suggestions, /api/crawl-status) → "
     "a React single-page presentation layer."),

    ("Cluster IDs are mapped to categories with a greedy one-to-one "
     "assignment: clusters are processed in order of how strongly a single "
     "category dominates them, and each claims its most common label if "
     "unclaimed — guaranteeing a valid bijective mapping, unlike "
     "independent per-cluster majority voting (which could, in principle, "
     "assign two clusters to the same category). Labels are used only for "
     "this post-hoc mapping, never as training signal.",
     "Cluster IDs are mapped to categories with a greedy one-to-one "
     "assignment: clusters are processed by how strongly one category "
     "dominates them, and each claims its top label if unclaimed — "
     "guaranteeing a valid bijective mapping, unlike independent "
     "per-cluster majority voting. Labels are used only for this post-hoc "
     "mapping, never as training signal."),

    ("2 of 3 live classification tests produced the expected category (a "
     "real 'Entertainment' sentence was classified as 'Economics'), using "
     "the genuinely trained K-Means model — no keyword override is "
     "applied (Appendix B). The misclassification is honest evidence of "
     "the same vocabulary overlap discussed above, not a fabricated "
     "failure: short, topic-light sentences are hardest for a "
     "distance-based unsupervised model.",
     "2 of 3 live classification tests produced the expected category (a "
     "real 'Entertainment' sentence was classified as 'Economics') using "
     "the genuinely trained model — no keyword override is applied "
     "(Appendix B). This reflects the same vocabulary overlap discussed "
     "above, not a fabricated failure."),

    ("This project implements two complementary IR approaches in one "
     "application. Task 1 is a complete vertical search engine for the "
     "Coventry PurePortal, using a polite, robots.txt-compliant crawler, "
     "TF-IDF Vector Space Model, and cosine similarity ranking to return "
     "the top-K=10 results for arbitrary queries. The 90-day crawl "
     "schedule, enforced robots.txt compliance, and MongoDB architecture "
     "fulfil the specification.",
     "This project implements two complementary IR approaches. Task 1 is "
     "a complete vertical search engine for the Coventry PurePortal, using "
     "a polite, robots.txt-compliant crawler, TF-IDF Vector Space Model, "
     "and cosine similarity ranking to return the top-K=10 results for "
     "arbitrary queries — fulfilling the specification's 90-day schedule, "
     "robots.txt, and MongoDB requirements."),

    ("A vertical search engine focuses retrieval on a specific domain "
     "rather than the entire web — here, the Coventry University Centre "
     "for Healthcare and Community Transformation PurePortal, a repository "
     "of research outputs and researcher profiles. The problem is locating "
     "specific publications or researchers via natural language queries "
     "within a portal that exposes no public search API.",
     "A vertical search engine focuses retrieval on one domain rather "
     "than the entire web — here, the Coventry Centre for Healthcare and "
     "Community Transformation PurePortal. The problem: locating "
     "publications or researchers via natural-language queries within a "
     "portal with no public search API."),

    ("The dataset is imbalanced (Economics 72, Entertainment 42, Politics "
     "84) because a live RSS feed exposes only a limited number of recent "
     "items — Entertainment's feed yields fewer — and the Wikipedia "
     "top-up is itself rate-limited. As the literature predicts, K-Means "
     "centroids drift toward the majority category under imbalance, "
     "visible in the confusion matrix.",
     "The dataset is imbalanced (Economics 72, Entertainment 42, Politics "
     "84) because a live RSS feed exposes only limited recent items — "
     "Entertainment's feed yields fewer — and the Wikipedia top-up is "
     "itself rate-limited. As the literature predicts, K-Means centroids "
     "drift toward the majority category under imbalance (visible above)."),

    ("Task 2 uses three collections. news_documents stores each article's "
     "title, url, content, source, category, cleaned_text, fingerprint, "
     "cluster_id/label, and pca_x/y. news_model_runs stores the complete "
     "K-Means state (serialised vectoriser, SVD reducer, normaliser, and "
     "KMeans objects), the cluster_map, evaluation metrics, and training "
     "timestamp. news_classifications logs each user classification "
     "request with its input, prediction, confidence, and timestamp.",
     "Task 2 uses three collections: news_documents (article text, "
     "category, cleaned_text, cluster_id/label, pca_x/y); news_model_runs "
     "(serialised vectoriser/SVD/normaliser/KMeans, cluster_map, "
     "evaluation metrics, training timestamp); and news_classifications "
     "(each user request's input, prediction, confidence, timestamp)."),

    ("Polite crawling is implemented throughout: a five-second delay "
     "between requests (matching PurePortal's own robots.txt Crawl-Delay), "
     "a descriptive User-Agent, URL deduplication, and GET requests only. "
     "robots.txt compliance is actually enforced — can_fetch() is called "
     "from fetch_page() before every request, not merely demonstrated in "
     "a comment (see Section 2.4.1 for a real bug this uncovered).",
     "Polite crawling: a five-second delay matching PurePortal's own "
     "robots.txt Crawl-Delay, a descriptive User-Agent, URL "
     "deduplication, and GET requests only. robots.txt compliance is "
     "actually enforced — can_fetch() runs before every request, not just "
     "commented (Section 2.4.1 describes a real bug this uncovered)."),

    ("As an automated, reproducible relevance proxy (rather than "
     "subjective manual judgement), Table 2 reports lexical-overlap "
     "precision: the fraction of top-10 results whose title contains a "
     "query word. This is a conservative lower bound, since a genuinely "
     "relevant document can rank highly on body-text similarity alone "
     "without repeating query terms in its title.",
     "As an automated relevance proxy, Table 2 reports lexical-overlap "
     "precision: the fraction of top-10 results whose title contains a "
     "query word — a conservative lower bound, since a relevant document "
     "can rank highly on body-text similarity alone."),

    ("Tests T1–T4 demonstrate correct operation: the system ranks the "
     "most relevant document highest for T1 ('...Mental Well-Being', "
     "score 0.2982), and correctly surfaces the queried researcher's own "
     "publications for author-name queries T2/T3. T4 returns a full page "
     "for a multi-term domain query, confirming keyword search works "
     "alongside author and title search.",
     "T1–T4 demonstrate correct operation: the system ranks the most "
     "relevant document highest for T1 (score 0.2982), correctly surfaces "
     "the researcher's own publications for T2/T3, and T4 returns a full "
     "page for a multi-term query — keyword search works alongside author "
     "and title search."),

    ("Third, the most recent full crawl visited 282 pages with 78 fetch "
     "errors (timeouts/transient HTTP errors on a minority of pages) — "
     "normal for a large, polite crawl, and it did not prevent the corpus "
     "being built. More retries and a longer timeout would recover a few "
     "more pages.",
     "Third, the most recent full crawl visited 282 pages with 78 fetch "
     "errors (timeouts/transient HTTP errors on a minority) — normal for "
     "a large polite crawl, and it did not prevent the corpus being "
     "built."),

    ("Evaluation shows the VSM search correctly ranks the most relevant "
     "publication first for 'mental health'. The one live classification "
     "miss reflects genuine Economics/Entertainment/Politics vocabulary "
     "overlap, not a system fault. The platform provides a professional, "
     "responsive interface with relevance-score visualisations, donut "
     "charts, PCA scatter plots, and real-time text classification.",
     "Evaluation shows the VSM search correctly ranks the most relevant "
     "publication first for 'mental health'; the one live classification "
     "miss reflects genuine vocabulary overlap, not a system fault. The "
     "platform provides a responsive interface with relevance-score "
     "visualisations, donut charts, PCA scatter plots, and real-time "
     "classification."),

    ("Several limitations remain. First, the corpus (77 indexed "
     "publications, from 83 crawled pages) is modest next to a production "
     "PurePortal deployment covering hundreds of publications; some "
     "generic terms carry limited discriminative IDF weight simply "
     "because the corpus is small — a larger corpus would improve IDF "
     "calibration.",
     "Several limitations remain. First, the corpus (77 indexed "
     "publications) is modest next to a production deployment covering "
     "hundreds; some generic terms carry limited discriminative IDF "
     "weight simply because the corpus is small."),

    ("The system is implemented using Python (Flask) for the backend API, "
     "React (Vite) for the frontend, MongoDB Atlas for storage, and "
     "NLTK/scikit-learn for text processing. The crawler uses requests and "
     "BeautifulSoup, performing a breadth-first traversal from the seed "
     "URL that follows only Research Output and Profile links.",
     "The system uses Python (Flask) for the backend API, React (Vite) "
     "for the frontend, MongoDB Atlas for storage, and NLTK/scikit-learn "
     "for text processing. The crawler (requests + BeautifulSoup) "
     "performs a breadth-first traversal from the seed URL, following "
     "only Research Output and Profile links."),

    ("T5 ('healthcare community transformation') now returns 54 results, "
     "with the Centre's own organisation/network pages ranked highest "
     "(score 0.6852) — these terms are no longer stopwords, so the query "
     "behaves as expected. T6, an out-of-domain query, correctly returns "
     "zero results, confirming the system does not fabricate matches.",
     "T5 ('healthcare community transformation') now returns 54 results, "
     "topped by the Centre's own organisation/network pages (score "
     "0.6852). T6, an out-of-domain query, correctly returns zero "
     "results — the system does not fabricate matches."),

    ("Cluster visualisation applies PCA to reduce the TF-IDF matrix to "
     "two principal components — the directions of maximum, then maximum "
     "remaining, variance. Each document is plotted as a point coloured "
     "by its assigned cluster; well-separated clusters indicate strong "
     "vocabulary distinctions, while overlap indicates shared "
     "terminology.",
     "Cluster visualisation applies PCA to reduce the TF-IDF matrix to "
     "two principal components. Each document is plotted as a point "
     "coloured by its assigned cluster; separation indicates strong "
     "vocabulary distinctions, overlap indicates shared terminology."),

    ("Tasks 1 and 2 represent complementary IR paradigms. Task 1 is "
     "precision-oriented: given a query, it returns the most relevant "
     "documents from a curated corpus via principled ranking. Task 2 is "
     "discovery-oriented: given a growing collection, it organises "
     "documents into meaningful categories without explicit queries.",
     "Tasks 1 and 2 are complementary IR paradigms: Task 1 is "
     "precision-oriented, returning the most relevant documents for a "
     "query via principled ranking; Task 2 is discovery-oriented, "
     "organising a growing collection into categories without explicit "
     "queries."),

    ("Scalability differs: VSM search is O(N×V) per query — trivial at 77 "
     "documents, but an inverted index with approximate nearest-neighbour "
     "search would be needed at scale. K-Means scales better in practice, "
     "since clustering runs offline and only centroid distances are "
     "needed at classification time.",
     "Scalability differs: VSM search is O(N×V) per query — trivial at 77 "
     "documents, but needing an inverted index at scale. K-Means scales "
     "better in practice, since clustering runs offline and only centroid "
     "distances are needed at classification time."),

    ("Where t is a term, d a document, count(t,d) its frequency in d, |d| "
     "the total terms in d, N the corpus size (83), and df(t) the number "
     "of documents containing t. Document vectors are L2-normalised, so "
     "document length does not bias similarity scores.",
     "Where t is a term, d a document, |d| the total terms in d, N the "
     "corpus size (83), and df(t) the documents containing t. Document "
     "vectors are L2-normalised, so length does not bias similarity."),

    ("The challenge is high dimensionality: a TF-IDF matrix for 198 "
     "documents may contain thousands of features, and K-Means must "
     "partition this space into three coherent, meaningful clusters. The "
     "system also provides real-time classification of new user-supplied "
     "text, demonstrating the trained model's operational utility.",
     "The challenge is high dimensionality: a TF-IDF matrix for 198 "
     "documents may contain thousands of features, which K-Means must "
     "partition into three coherent clusters. The system also classifies "
     "new user-supplied text in real time."),

    ("Crawler scheduling is managed by scheduler.py (start_scheduler()), "
     "using APScheduler's BlockingScheduler with an IntervalTrigger set to "
     "days=CRAWL_INTERVAL_DAYS (90 by default) — implementing the strict "
     "3-month interval required by the specification. The scheduler runs "
     "the crawl immediately on startup, then re-runs every 3 months:",
     "Crawler scheduling (scheduler.py: start_scheduler()) uses "
     "APScheduler's BlockingScheduler with an IntervalTrigger set to "
     "days=CRAWL_INTERVAL_DAYS (90 by default), running immediately on "
     "startup and then every 3 months:"),

    ("Both are unsupervised in their core mechanism — the VSM ranks purely "
     "on term-distribution statistics, and K-Means clusters purely on "
     "distance, using labels only post-hoc for interpretation. Neither "
     "requires labelled training data, so both apply to new domains "
     "without annotation effort.",
     "Both are unsupervised at their core — VSM ranks purely on term "
     "distribution, K-Means clusters purely on distance, using labels "
     "only post-hoc — so both apply to new domains without annotation."),

    ("TF-IDF was chosen over plain bag-of-words because it down-weights "
     "high-frequency, low-information terms, improving cluster separation "
     "(Salton & McGill, 1983; Steinbach, Karypis, & Kumar, 2000). "
     "k-means++ initialisation (Arthur & Vassilvitskii, 2007) spreads "
     "initial centroids probabilistically, improving convergence over "
     "random initialisation.",
     "TF-IDF was chosen over bag-of-words as it down-weights "
     "high-frequency, low-information terms, improving separation (Salton "
     "& McGill, 1983; Steinbach et al., 2000). k-means++ (Arthur & "
     "Vassilvitskii, 2007) spreads initial centroids probabilistically, "
     "improving convergence over random initialisation."),
]

ok = 0
for old, new in PAIRS:
    if replace_exact(doc, old, new):
        ok += 1
print(f"Replaced {ok}/{len(PAIRS)} paragraphs.")

# Cut the "four query types" paragraph — redundant with the query
# processing explanation already given.
delete_exact(doc,
    "The system supports four query types: publication title queries; "
    "author name queries, matching name tokens against profile and "
    "publication documents; keyword queries for domain concepts; and "
    "multi-term phrase queries, where multiple matches produce additive "
    "cosine contributions.",
    label="four query types paragraph")

doc.save(REPORT_PATH)
print("Saved.")
