import os
from docx import Document

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")


def set_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


doc = Document(REPORT_PATH)
paras = doc.paragraphs
idx = None
for i, p in enumerate(paras):
    if p.text.strip().endswith("Key elements include:"):
        idx = i
        break

if idx is None:
    print("anchor not found")
else:
    set_paragraph_text(paras[idx],
        "The search interface is a React single-page application "
        "(SearchPage.jsx), styled after Google Scholar's functional "
        "simplicity with a dark academic aesthetic. Key elements include a "
        "search input with autocomplete suggestions and quick-term chips; "
        "an index statistics strip (live document/term counts, crawl "
        "schedule); result cards with clickable title, clickable author "
        "profile links, publication date, and a visual relevance bar; and "
        "loading, empty, and error states with functional pagination.")
    to_remove = paras[idx + 1: idx + 9]
    for p in to_remove:
        p._element.getparent().remove(p._element)
    print(f"Merged and removed {len(to_remove)} bullets.")

doc.save(REPORT_PATH)
print("Saved.")
