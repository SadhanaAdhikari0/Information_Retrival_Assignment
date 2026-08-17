"""
trim_report2.py — Second, more aggressive trim pass: merges bulleted lists
into compact prose (deleting the now-redundant bullet paragraphs) and
shortens remaining verbose paragraphs, to bring the main body further
toward the ~2,000-word target.
"""
import os
from docx import Document

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")


def set_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def replace_exact(doc, old, new, label=""):
    for p in doc.paragraphs:
        if p.text.strip() == old.strip():
            set_paragraph_text(p, new)
            return True
    print(f"  [WARN] exact paragraph not found: {label or old[:60]}")
    return False


def merge_and_delete_following(doc, anchor_text, new_anchor_text, n_delete):
    """Find the paragraph matching anchor_text, replace its text with
    new_anchor_text, then delete the next n_delete paragraphs (the bullet
    items being folded into it)."""
    paras = doc.paragraphs
    idx = None
    for i, p in enumerate(paras):
        if p.text.strip() == anchor_text.strip():
            idx = i
            break
    if idx is None:
        print(f"  [WARN] anchor not found: {anchor_text[:60]}")
        return False

    set_paragraph_text(paras[idx], new_anchor_text)
    to_remove = paras[idx + 1: idx + 1 + n_delete]
    for p in to_remove:
        p._element.getparent().remove(p._element)
    return True


doc = Document(REPORT_PATH)

# ---- Merge bulleted lists into compact prose ----
MERGES = [
    ("For each research output, the following metadata is extracted:",
     "For each research output, the crawler extracts: title, author names "
     "and profile URLs, publication date, abstract, publication type, "
     "keywords, full page text (for indexing), source URL, and crawl/"
     "update timestamps.", 9),

    ("Key interface elements include:",
     "Key elements include a search input with autocomplete suggestions "
     "and quick-term chips; an index statistics strip (live document/term "
     "counts, crawl schedule); result cards with clickable title, "
     "clickable author profile links, publication date, and a visual "
     "relevance bar; and loading, empty, and error states with functional "
     "pagination.", 8),

    ("Collect a real, citable document set across Economics, Entertainment, "
     "and Politics from live RSS feeds and Wikipedia.",
     "Collect a real, citable document set across Economics, Entertainment, "
     "and Politics from live RSS feeds and Wikipedia; preprocess "
     "consistently and vectorise with TF-IDF (unigrams + bigrams); train "
     "K-Means (K=3, k-means++) and map clusters to categories via a greedy "
     "one-to-one assignment; reduce to 2D via PCA for visualisation and "
     "compute the silhouette score; allow users to classify arbitrary "
     "text, persisting results to MongoDB.", 8),

    ("Economics: ECONOMICS_RSS_URL (e.g., BBC Business RSS)",
     "Feed URLs are configured via environment variables — "
     "ECONOMICS_RSS_URL, ENTERTAINMENT_RSS_URL, POLITICS_RSS_URL — e.g. "
     "BBC Business/Entertainment/Politics RSS.", 2),

    ("HTML Cleaning: HTML tags and common HTML entities are stripped using "
     "regular expressions, as RSS content often contains embedded markup.",
     "The pipeline strips HTML tags/entities (RSS content often embeds "
     "markup), lowercases text, removes non-alphabetic characters, "
     "tokenises with NLTK, removes standard English stop-words plus "
     "news-domain terms ('said', 'would', 'could', 'reuters', 'bbc', "
     "'cnn', 'ap', 'afp'), and stems remaining tokens.", 5),
]
for anchor, new_text, n in MERGES:
    merge_and_delete_following(doc, anchor, new_text, n)
    print(f"Merged: {anchor[:50]}")

doc.save(REPORT_PATH)
print("Saved (merges).")
