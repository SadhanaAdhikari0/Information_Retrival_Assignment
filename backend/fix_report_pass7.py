"""fix_report_pass7.py — mop up remaining stale file/code references."""
import os
from docx import Document
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), ".env"))
from pymongo import MongoClient

db = MongoClient(os.environ["MONGODB_URI"])["vertical_search_engine"]
REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")
n_term_index = db.term_index.count_documents({})


def set_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def replace_in_all_text(doc, replacements):
    def process(paragraph):
        full = "".join(r.text for r in paragraph.runs)
        if not full:
            return
        new_full = full
        for old, new in replacements:
            new_full = new_full.replace(old, new)
        if new_full != full:
            set_paragraph_text(paragraph, new_full)
    for p in doc.paragraphs:
        process(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process(p)


FIXES = [
    ("The crawler (crawler.py) targets the Centre for Healthcare and "
     "Community Transformation organisation page at:",
     "The crawler (scheduler.py: crawl()) targets the Centre for Healthcare "
     "and Community Transformation organisation page at:"),

    (f"The term_index collection stores the pre-computed IDF value for each "
     f"of the 2,174 unique terms in the vocabulary. During indexing "
     f"(rebuild_index.py), document TF-IDF vectors are computed and stored "
     f"in the doc_vectors collection. During search, query vectors are "
     f"computed using the same stored IDF values, ensuring consistency "
     f"between document and query representations.",
     f"The term_index collection stores the pre-computed IDF value for each "
     f"of the {n_term_index:,} unique terms in the vocabulary. During "
     f"indexing (scheduler.py: build_index()), document TF-IDF vectors are "
     f"computed and stored in the doc_vectors collection. During search, "
     f"query vectors are computed using the same stored IDF values, "
     f"ensuring consistency between document and query representations."),

    ("A.5 Robots.txt Compliance (crawler/crawler.py — commented)",
     "A.5 Robots.txt Compliance (backend/rss_collector.py — actually "
     "enforced before every article fetch, not just commented)"),

    ('''# from urllib.robotparser import RobotFileParser
# def load_robots_txt(base_url: str):
#     parsed = urlparse(base_url)
#     robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
#     rp = RobotFileParser()
#     rp.set_url(robots_url)
#     try:
#         resp = requests.get(robots_url, headers={"User-Agent": USER_AGENT}, timeout=10)
#         if resp.status_code == 200:
#             rp.parse(resp.text.splitlines())
#         else:
#             rp = None
#     except Exception:
#         rp = None
#     return rp
#
# def is_allowed(rp, url: str) -> bool:
#     if rp is None:
#         return True
#     return rp.can_fetch("*", url)''',
     '''_ROBOTS_CACHE = {}

def check_robots_txt(url: str, user_agent: str = "*") -> bool:
    """Checks robots.txt before fetching \\u2014 called from
    fetch_rss_articles() for every article URL, not just commented out."""
    base_url = f"{urlparse(url).scheme}://{urlparse(url).netloc}"
    if base_url not in _ROBOTS_CACHE:
        try:
            rp = RobotFileParser()
            rp.set_url(f"{base_url}/robots.txt")
            rp.read()
            _ROBOTS_CACHE[base_url] = rp
        except Exception:
            _ROBOTS_CACHE[base_url] = None
    rp = _ROBOTS_CACHE[base_url]
    return True if rp is None else rp.can_fetch(user_agent, url)'''),

    ("E.1 Crawler Link Filtering (crawler/crawler.py)",
     "E.1 Crawler Link Filtering (backend/scheduler.py)"),

    ('''def extract_pureportal_links(html: str):
    soup = BeautifulSoup(html, 'html.parser')
    links = set()
    for a in soup.find_all('a', href=True):
        href = a['href']
        if '/publications/' in href or '/persons/' in href:
            links.add(href)
    return list(links)''',
     '''def is_relevant_url(url: str) -> bool:
    """Accept only Research Output and Profile URLs; rejects everything else."""
    parsed = urlparse(url)
    if not parsed.netloc.endswith(ALLOWED_DOMAIN):
        return False
    return any(parsed.path.startswith(p) for p in ALLOWED_PATH_PREFIXES)
    # ALLOWED_PATH_PREFIXES = ("/en/publications/", "/en/persons/",
    #                          "/en/research-output/",
    #                          "/en/organisations/centre-for-healthcare")'''),
]
doc = Document(REPORT_PATH)
replace_in_all_text(doc, FIXES)
doc.save(REPORT_PATH)
print("Saved pass 7.")
