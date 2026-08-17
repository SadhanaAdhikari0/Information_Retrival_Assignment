"""fix_report_pass8.py — final robots.txt honesty pass (real bug found+fixed)."""
import os
from docx import Document

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Documentation",
                            "ST7071CEM_Information_Retrieval_Report.docx")


def set_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def replace_in_all_text(doc, replacements):
    def process(paragraph):
        full = "".join(r.text for r in paragraph.runs)
        if not full:
            return
        new_full = full
        for old, new in replacements:
            new_full = new_full.replace(old, new)
        if new_full != full:
            set_paragraph_text(paragraph, new_full)
    for p in doc.paragraphs:
        process(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process(p)


FIXES = [
    ("Polite crawling practices are implemented throughout: a six-second "
     "delay between requests (CRAWL_DELAY = 6), a descriptive user-agent "
     "string, URL deduplication to prevent revisiting pages, and standard "
     "HTTP GET requests only. The crawler also includes commented "
     "implementation code demonstrating robots.txt compliance using "
     "Python's RobotFileParser, as required by the assignment "
     "specification. In production, this implementation would check "
     "is_allowed(rp, url) before fetching each page.",
     "Polite crawling practices are implemented throughout: a five-second "
     "delay between requests (CRAWL_DELAY_SECONDS = 5, matching "
     "PurePortal's own robots.txt Crawl-Delay directive), a descriptive "
     "User-Agent string, URL deduplication to prevent revisiting pages, and "
     "standard HTTP GET requests only. robots.txt compliance is actually "
     "enforced (can_fetch() is called from fetch_page() before every "
     "request), not merely demonstrated in a comment — see Section 2.4.1 "
     "for a bug this uncovered."),

    ("Ethical web crawling is a fundamental requirement of modern "
     "Information Retrieval systems. To ensure strict adherence to website "
     "policies, the crawler implements a robust robots.txt parser using "
     "Python's native urllib.robotparser module. Before executing any HTTP "
     "requests against a domain, the system resolves the base URL, fetches "
     "the robots.txt file, and verifies if the user-agent is explicitly "
     "permitted to scrape the target directory. This prevents the system "
     "from causing excessive server load on disallowed paths or violating "
     "the host's terms of service.",
     "Ethical web crawling is a fundamental requirement of modern "
     "Information Retrieval systems. The crawler uses Python's native "
     "urllib.robotparser to check every URL against PurePortal's robots.txt "
     "before it is fetched. Building this surfaced a genuine bug worth "
     "reporting: RobotFileParser.read() fetches robots.txt via bare "
     "urllib.request with no custom headers, and PurePortal's edge "
     "protection returns HTTP 403 to urllib's default User-Agent (confirmed "
     "independently: the identical URL returns 200 via the requests "
     "library). RobotFileParser silently treats a 403 response as "
     "\"disallow everything\", which meant the very first version of this "
     "check blocked the crawler from fetching even its own seed URL. The "
     "fix — fetch robots.txt text via requests with the same descriptive "
     "User-Agent used for every other request, then hand that text to "
     "RobotFileParser.parse() instead of .read() — resolved it. "
     "PurePortal's robots.txt (checked 2026-08-15) permits every path this "
     "crawler visits and specifies Crawl-Delay: 5, which the crawler's "
     "5-second delay matches."),

    ('''# from urllib.robotparser import RobotFileParser
# def load_robots_txt(base_url: str):
#     parsed = urlparse(base_url)
#     robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
#     rp = RobotFileParser()
#     rp.set_url(robots_url)
#     try:
#         resp = requests.get(robots_url, headers={"User-Agent": USER_AGENT}, timeout=10)
#         if resp.status_code == 200:
#             rp.parse(resp.text.splitlines())
#         else:
#             rp = None
#     except Exception:
#         rp = None
#     return rp
#
# def is_allowed(rp, url: str) -> bool:
#     if rp is None:
#         return True
#     return rp.can_fetch("*", url)''',
     '''# Fetch robots.txt via requests (NOT RobotFileParser.read(), which uses
# bare urllib.request and gets HTTP 403 from PurePortal's edge protection,
# silently causing RobotFileParser to disallow-everything). Parse the
# fetched text with RobotFileParser.parse() instead.

def _get_robots_parser():
    global _robots_parser
    if _robots_parser is None:
        _robots_parser = RobotFileParser()
        try:
            resp = requests.get(f"https://{ALLOWED_DOMAIN}/robots.txt",
                                 headers={"User-Agent": USER_AGENT}, timeout=10)
            resp.raise_for_status()
            _robots_parser.parse(resp.text.splitlines())
        except Exception as e:
            print(f"  [robots.txt] fetch failed: {e} \\u2014 treating as permissive")
            _robots_parser = False
    return _robots_parser or None

def can_fetch(url: str, user_agent: str = "ST7071CEM-IR-Bot") -> bool:
    rp = _get_robots_parser()
    return True if rp is None else rp.can_fetch(user_agent, url)

# Wired into fetch_page():
#   if not can_fetch(url, user_agent=USER_AGENT):
#       print(f"  [robots.txt] Disallowed, skipping: {url}")
#       return None'''),
]
doc = Document(REPORT_PATH)
replace_in_all_text(doc, FIXES)
doc.save(REPORT_PATH)
print("Saved pass 8.")
